from tkinter import *
import docx

class mc:
	def __init__(self,q,a,b,c,d):
		self.q = q
		self.a = a
		self.b = b
		self.c = c 
		self.d = d
doc = docx.Document("Exercise6a.docx")

totpara = len(doc.paragraphs)

for x in range(0,totpara):

	print(doc.paragraphs[x].text)
	if 'True/False' in doc.paragraphs[x].text or 'State True/false' in doc.paragraphs[x].text:
		break
	# if  doc.paragraphs[x].runs: 
	# 	if doc.paragraphs[x].runs[0].font.highlight_color: print('ANSWER',end='')

print('START TRUE AND FALSE SECTION')

for y in range(x+1,totpara):
	if 'TRUE' in doc.paragraphs[y].text or 'FALSE' in doc.paragraphs[y].text:
		continue
	print(doc.paragraphs[y].text)

	

